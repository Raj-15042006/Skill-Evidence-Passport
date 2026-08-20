import io
import os
import re
import logging
from typing import Union, Optional
from PIL import Image
import pypdf
import docx

try:
    import pytesseract
    HAS_PYTESSERACT = True
except ImportError:
    HAS_PYTESSERACT = False

from app.models.schemas import ExtractionResult

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    """Normalizes whitespace and strips control characters."""
    if not text:
        return ""
    # Normalize multiple whitespace/newlines to single space/newlines
    text = re.sub(r'[\r\n]+', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()


def extract_from_pdf(stream_or_path: Union[io.BytesIO, str]) -> ExtractionResult:
    try:
        reader = pypdf.PdfReader(stream_or_path)
        extracted_pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                extracted_pages.append(t)
        combined_text = clean_text("\n".join(extracted_pages))
        low_confidence = len(combined_text) < 20
        return ExtractionResult(
            text=combined_text,
            low_confidence_extraction=low_confidence,
            extraction_method="pdf"
        )
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        return ExtractionResult(
            text="",
            low_confidence_extraction=True,
            extraction_method="pdf",
            error_message=str(e)
        )


def extract_from_docx(stream_or_path: Union[io.BytesIO, str]) -> ExtractionResult:
    try:
        doc = docx.Document(stream_or_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())
        combined_text = clean_text("\n".join(full_text))
        low_confidence = len(combined_text) < 15
        return ExtractionResult(
            text=combined_text,
            low_confidence_extraction=low_confidence,
            extraction_method="docx"
        )
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return ExtractionResult(
            text="",
            low_confidence_extraction=True,
            extraction_method="docx",
            error_message=str(e)
        )


def extract_from_image(stream_or_path: Union[io.BytesIO, str, Image.Image]) -> ExtractionResult:
    low_confidence = False
    extracted_text = ""
    error_msg = None

    try:
        if isinstance(stream_or_path, Image.Image):
            img = stream_or_path
        else:
            img = Image.open(stream_or_path)

        if not HAS_PYTESSERACT:
            return ExtractionResult(
                text="",
                low_confidence_extraction=True,
                extraction_method="ocr_unavailable",
                error_message="pytesseract library not installed"
            )

        try:
            # Use image_to_data to inspect confidence scores
            data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            confidences = []
            words = []
            for i in range(len(data['text'])):
                word = data['text'][i].strip()
                conf = int(data['conf'][i])
                if word:
                    words.append(word)
                    if conf >= 0:  # -1 indicates unassigned confidence
                        confidences.append(conf)

            extracted_text = clean_text(" ".join(words))
            avg_conf = (sum(confidences) / len(confidences)) if confidences else 0

            if avg_conf < 50 or len(words) < 3:
                low_confidence = True
        except Exception as ocr_err:
            logger.warning(f"Tesseract OCR execution failed/unavailable: {ocr_err}. Falling back gracefully.")
            error_msg = str(ocr_err)
            low_confidence = True
            # Try simple image string fallback if image_to_data fails
            try:
                extracted_text = clean_text(pytesseract.image_to_string(img))
            except Exception:
                extracted_text = ""

    except Exception as e:
        logger.error(f"Error processing image file: {e}")
        return ExtractionResult(
            text="",
            low_confidence_extraction=True,
            extraction_method="image_ocr",
            error_message=str(e)
        )

    return ExtractionResult(
        text=extracted_text,
        low_confidence_extraction=low_confidence,
        extraction_method="image_ocr",
        error_message=error_msg
    )


def extract_text(
    source: Union[str, bytes],
    file_type: Optional[str] = None
) -> ExtractionResult:
    """
    Extracts plain text from raw bytes or string source (file path, URL, or plain text).
    Supports PDF, DOCX, images (OCR), and plain text.
    """
    if not source:
        return ExtractionResult(
            text="",
            low_confidence_extraction=True,
            extraction_method="empty_input"
        )

    # 1. Handle byte input
    if isinstance(source, bytes):
        stream = io.BytesIO(source)
        
        # Check explicit file_type or detect from header bytes
        fmt = (file_type or "").lower().strip(".")
        if fmt == "pdf" or source.startswith(b"%PDF"):
            return extract_from_pdf(stream)
        elif fmt in ("docx", "doc") or source.startswith(b"PK"):
            return extract_from_docx(stream)
        elif fmt in ("png", "jpg", "jpeg", "bmp", "tiff"):
            return extract_from_image(stream)
        else:
            # Try image decode first
            try:
                img = Image.open(stream)
                return extract_from_image(img)
            except Exception:
                pass
            
            # Fallback to text decoding
            try:
                decoded = source.decode("utf-8", errors="ignore")
                cleaned = clean_text(decoded)
                return ExtractionResult(
                    text=cleaned,
                    low_confidence_extraction=len(cleaned) < 5,
                    extraction_method="plain_text_bytes"
                )
            except Exception as e:
                return ExtractionResult(
                    text="",
                    low_confidence_extraction=True,
                    extraction_method="raw_bytes",
                    error_message=str(e)
                )

    # 2. Handle string input
    if isinstance(source, str):
        # If string is a path to an existing file
        if os.path.exists(source) and os.path.isfile(source):
            ext = os.path.splitext(source)[1].lower().strip(".")
            if ext == "pdf":
                return extract_from_pdf(source)
            elif ext in ("docx", "doc"):
                return extract_from_docx(source)
            elif ext in ("png", "jpg", "jpeg", "bmp", "tiff"):
                return extract_from_image(source)
            else:
                try:
                    with open(source, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    cleaned = clean_text(text)
                    return ExtractionResult(
                        text=cleaned,
                        low_confidence_extraction=len(cleaned) < 5,
                        extraction_method="file_path_text"
                    )
                except Exception as e:
                    return ExtractionResult(
                        text="",
                        low_confidence_extraction=True,
                        extraction_method="file_path",
                        error_message=str(e)
                    )

        # Plain text or URL pass-through
        cleaned = clean_text(source)
        return ExtractionResult(
            text=cleaned,
            low_confidence_extraction=len(cleaned) < 5,
            extraction_method="plain_text_string"
        )
