import os
import io
from PIL import Image, ImageDraw, ImageFont
import docx
import pypdf

def create_fixtures():
    fixtures_dir = os.path.dirname(__file__)
    os.makedirs(fixtures_dir, exist_ok=True)

    # 1. Plain text fixture
    txt_path = os.path.join(fixtures_dir, "sample_text.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("Developed a Python web application with FastAPI and Pytest. Configured automated CI/CD pipeline.")

    # 2. DOCX fixture
    docx_path = os.path.join(fixtures_dir, "sample_project.docx")
    doc = docx.Document()
    doc.add_heading("Python Skill Evidence Report", level=1)
    doc.add_paragraph("Implemented asynchronous REST endpoints and unit tests for data analysis.")
    doc.save(docx_path)

    # 3. Image fixture (PNG)
    img_path = os.path.join(fixtures_dir, "sample_scan.png")
    img = Image.new('RGB', (400, 100), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((10, 40), "Python Certification Document", fill=(0, 0, 0))
    img.save(img_path)

    # 4. PDF fixture
    pdf_path = os.path.join(fixtures_dir, "sample_resume.pdf")
    writer = pypdf.PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    # Simple blank PDF page for extraction testing
    with open(pdf_path, "wb") as f:
        writer.write(f)

    print("Test fixtures generated successfully.")

if __name__ == "__main__":
    create_fixtures()
